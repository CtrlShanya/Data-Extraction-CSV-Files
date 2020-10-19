import pandas as pd
import time
from docx import Document

#chcp 65001
#set PYTHONIOENCODING=utf-8

#pip install python-docx


df_R = pd.read_csv("Rashi.csv", encoding="ISO-8859-1")

code = df_R["Rashi"]
one = df_R["Rashi / Yog"]
two = df_R["Auspicious"]
three = df_R["Inauspicious"]
four = df_R["Badhak /Marak"]
five = df_R["BVR Marak"]
six = df_R["Marak"]
seven = df_R["8th"]
#8th Lord, if not related to Trikona Lord can destroy the house where he is placed.
eight = df_R["2nd"]
#2nd & 12th lord Can be Karak  if 
nine = df_R["RY"]
#Raj Yog karak (RYK) & other Chap 5 P 132
ten = df_R["H1"]
eleven = df_R["Mental Tendencies"]
twelve = df_R["Physical Tendencies"]
thirteen = df_R["General Tendencies"]
document = Document()

c_ode = input('Enter R: ')
answer = input('Are you sure? ')
if(answer=='yes' or answer=='Yes'):
	for i in range(len(code)):
		if(c_ode==code[i]):
			document.add_heading(code[i], level=0)
			document.add_heading('Rashi/Yog', level=2)
			try:
				paragraph = document.add_paragraph(one[i])
				print(one[i])
			except:
				paragraph = document.add_paragraph('NULL')
			document.add_heading('Auspicious', level=2)
			try:
				paragraph = document.add_paragraph(two[i])
				print(two[i])
			except:
				paragraph = document.add_paragraph('NULL')
			document.add_heading('Inauspicious', level=2)
			try:
				paragraph = document.add_paragraph(three[i])
				print(three[i])
			except:
				paragraph = document.add_paragraph('NULL')
			document.add_heading('Badhak/Marak', level=2)
			try:
				paragraph = document.add_paragraph(four[i])
				print(four[i])
			except:
				paragraph = document.add_paragraph('NULL')
			document.add_heading('BVR Marak', level=2)
			try:
				paragraph = document.add_paragraph(five[i])
				print(five[i])
			except:
				paragraph = document.add_paragraph('NULL')
			document.add_heading('Marak', level=2)
			try:
				paragraph = document.add_paragraph(six[i])
				print(six[i])
			except:
				paragraph = document.add_paragraph('NULL')
			document.add_heading('8th Lord, if not related to Trikona Lord can destroy the house where he is placed', level=2)
			try:
				paragraph = document.add_paragraph(seven[i])
				print(seven[i])
			except:
				paragraph = document.add_paragraph('NULL')
			document.add_heading('2nd & 12th lord Can be Karak  if', level=2)
			try:
				paragraph = document.add_paragraph(eight[i])
				print(eight[i])
			except:
				paragraph = document.add_paragraph('NULL')
			document.add_heading('Raj Yog karak (RYK) & other Chap 5 P 132', level=2)
			try:
				paragraph = document.add_paragraph(nine[i])
				print(nine[i])
			except:
				paragraph = document.add_paragraph('NULL')
			document.add_heading('H1', level=2)
			try:
				paragraph = document.add_paragraph(ten[i])
				print(ten[i])
			except:
				paragraph = document.add_paragraph('NULL')
			document.add_heading('Mental Tendencies', level=2)
			try:
				paragraph = document.add_paragraph(eleven[i])
				print(eleven[i])
			except:
				paragraph = document.add_paragraph('NULL')
			document.add_heading('Physical Tendencies', level=2)
			try:
				paragraph = document.add_paragraph(twelve[i])
				print(twelve[i])
			except:
				paragraph = document.add_paragraph('NULL')
			document.add_heading('General Tendencies', level=2)
			try:
				paragraph = document.add_paragraph(thirteen[i])
				print(thirteen[i])
			except:
				paragraph = document.add_paragraph('NULL')
			
df_HPH = pd.read_csv("H-P-H.csv", encoding="ISO-8859-1")

code2 = df_HPH["Code"]
one2 = df_HPH["One"]
two2 = df_HPH["Two"]
three2 = df_HPH["Three"]
four2 = df_HPH["Four"]
five2 = df_HPH["Five"]


for j in range(0,100):
	c_ode2 = input('Enter H-P or H-H or other code: ')
	answer2 = input('Are you sure? ')
	if(answer2[0]=='y' or answer2[0]=='Y'):
		print(c_ode2[0]+" "+c_ode2[1]+" "+c_ode2[2]+" "+c_ode2[3])
		print(len(c_ode2))
		for i in range(len(code2)):
			print(code2[i])
			print(len(code2[i]))
			if(len(c_ode2)<8 and c_ode2==code2[i]):
				print("IN FIRST CONDITION CUZ LEN IS:")
				print(len(c_ode2))
				if(c_ode2[2]=='L' and c_ode2[0]=='H' and c_ode2[1]=='3' and c_ode2[3]=='P'):
					document.add_heading(code2[i], level=0)
					document.add_heading('H3 Lord with Pn', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Effect on courage', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('In debilitation and has conjunction with malefics person becomes unskilful and timid.', level=2)
					try:
						paragraph = document.add_paragraph(three2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[1]=='1' and c_ode2[2]=='P'):
					document.add_heading(code2[i], level=0)
					document.add_heading('H1 having', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Effect', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					try:
						paragraph = document.add_paragraph(three2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('If afflicted', level=2)
					try:
						paragraph = document.add_paragraph(four2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[1]=='2' and c_ode2[2]=='P'):
					document.add_heading(code2[i], level=0)
					document.add_heading('H2 having', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Effect', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[1]=='3' and c_ode2[2]=='P'):
					document.add_heading(code2[i], level=0)
					document.add_heading('H3 Lord with Pn', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Effect', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[1]=='4' and c_ode2[2]=='P'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Pn', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Effect', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('If Afflicted', level=2)
					try:
						paragraph = document.add_paragraph(three2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Career', level=2)
					try:
						paragraph = document.add_paragraph(four2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Signifactor', level=2)
					try:
						paragraph = document.add_paragraph(five2[i])
					except:
						paragraph = document.add_paragraph('NULL')				
				if(c_ode2[1]=='5' and c_ode2[2]=='P'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Pn in H5', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Effect', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('If Afflicted by Pn Effect', level=2)
					try:
						paragraph = document.add_paragraph(three2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					try:
						paragraph = document.add_paragraph(four2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[1]=='6' and c_ode2[2]=='P'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Pn', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Effect', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[1]=='7' and c_ode2[2]=='P'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Pn ', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Effect', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('If Afflicted by Pn Effect', level=2)
					try:
						paragraph = document.add_paragraph(three2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[1]=='8' and c_ode2[2]=='P'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Pn ', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Effect', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[1]=='9' and c_ode2[2]=='P'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Pn ', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Effect', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('If Afflicted by Pn Effect', level=2)
					try:
						paragraph = document.add_paragraph(three2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Results during dasha of Pn related to H9', level=2)
					try:
						paragraph = document.add_paragraph(four2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[1]=='1' and c_ode2[2]=='0' and c_ode2[3]=='P'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Pn ', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Effect', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('If Afflicted by Pn Effect', level=2)
					try:
						paragraph = document.add_paragraph(three2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Pn in H10 Profession', level=2)
					try:
						paragraph = document.add_paragraph(four2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[1]=='1' and c_ode2[2]=='1' and c_ode2[3]=='P'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Pn ', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('P 372', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('P 374', level=2)
					try:
						paragraph = document.add_paragraph(three2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[1]=='1' and c_ode2[2]=='2' and c_ode2[3]=='P'):
					document.add_heading(code2[i], level=0)
					document.add_heading('H12', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Wealth is spent on', level=2)
					try:
						paragraph = document.add_paragraph(three2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Planet connected with H12 Dasha Results', level=2)
					try:
						paragraph = document.add_paragraph(four2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[0]=='H' and c_ode2[2]=='L' and c_ode2[1]=='1'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Results', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[0]=='H' and c_ode2[2]=='L' and c_ode2[1]=='2'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Effect when H2L in Hn', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Effect when H2L in Hn (P 102)', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[0]=='H' and c_ode2[2]=='L' and c_ode2[1]=='3'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Effect when H3L in Hn', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[0]=='H' and c_ode2[2]=='L' and c_ode2[1]=='4'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Effect when H4L in Hn', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[0]=='H' and c_ode2[2]=='L' and c_ode2[1]=='5'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Effect when H5L in Hn', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[0]=='H' and c_ode2[2]=='L' and c_ode2[1]=='6'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Effect when H6L in Hn', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[0]=='H' and c_ode2[2]=='L' and c_ode2[1]=='7'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Results', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Affliction', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[0]=='H' and c_ode2[2]=='L' and c_ode2[1]=='8'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Results', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Affliction', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[0]=='H' and c_ode2[2]=='L' and c_ode2[1]=='9'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Results', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Affliction', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[0]=='H' and c_ode2[2]=='0' and c_ode2[1]=='1' and c_ode2[3]=='L'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Results', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Affliction', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[0]=='H' and c_ode2[2]=='1' and c_ode2[1]=='1' and c_ode2[3]=='L'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Results', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Affliction', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
				if(c_ode2[0]=='H' and c_ode2[2]=='2' and c_ode2[1]=='1' and c_ode2[3]=='L'):
					document.add_heading(code2[i], level=0)
					document.add_heading('Results', level=2)
					try:
						paragraph = document.add_paragraph(one2[i])
					except:
						paragraph = document.add_paragraph('NULL')
					document.add_heading('Affliction', level=2)
					try:
						paragraph = document.add_paragraph(two2[i])
					except:
						paragraph = document.add_paragraph('NULL')
			if(len(c_ode2)==8):
				if(c_ode2==code2[i]):
					print("IN CONDITION TWO CUZ LEN IS:")
					print(len(c_ode2))
					if(c_ode2[1]=='1'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H1Lord with HnL in Hn', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						document.add_heading('Effect when H1L in Hn with HnLord', level=2)
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='2'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H2Lord with HnL in Hn', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						document.add_heading('Effect when H2L in Hn with HnLord', level=2)
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='3'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H3Lord with HnL in Hn', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						document.add_heading('Effect when H3L in Hn with HnLord', level=2)
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='4'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H4Lord with HnL in Hn', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						document.add_heading('Effect when H4L in Hn with HnLord', level=2)
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='5'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H5Lord with HnL in Hn', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						document.add_heading('Effect when H5L in Hn with HnLord', level=2)
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='6'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H6Lord with HnL in Hn', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						document.add_heading('Effect when H6L in Hn with HnLord', level=2)
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='7'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H7Lord with HnL in Hn', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(three2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='8'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H8Lord', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						document.add_heading('Results', level=2)
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='9'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H9Lord', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						document.add_heading('With H Lord in same house during his dasha', level=2)
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(three2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='1' and c_ode2[2]=='0'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H10Lord', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(three2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='1' and c_ode2[2]=='1'):
						print("11")
						print("code2[i] is")
						print(code2[i])
						document.add_heading(code2[i], level=0)
						document.add_heading('H11Lord', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(three2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='1' and c_ode2[2]=='2'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H12Lord', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(three2[i])
						except:
							paragraph = document.add_paragraph('NULL')
			if(len(c_ode2)>8):
				if(c_ode2==code2[i]):
					print("IN CONDITION THREE CUZ LEN IS:")
					print(len(c_ode2))
					if(c_ode2[1]=='1'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H1Lord with HnL in Hn', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						document.add_heading('Effect when H1L in Hn with HnLord', level=2)
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='2'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H2Lord with HnL in Hn', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						document.add_heading('Effect when H2L in Hn with HnLord', level=2)
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='3'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H3Lord with HnL in Hn', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						document.add_heading('Effect when H3L in Hn with HnLord', level=2)
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='4'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H4Lord with HnL in Hn', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						document.add_heading('Effect when H4L in Hn with HnLord', level=2)
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='5'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H5Lord with HnL in Hn', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						document.add_heading('Effect when H5L in Hn with HnLord', level=2)
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='6'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H6Lord with HnL in Hn', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						document.add_heading('Effect when H6L in Hn with HnLord', level=2)
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='7'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H7Lord with HnL in Hn', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(three2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='8'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H8Lord', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						document.add_heading('Results', level=2)
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='9'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H9Lord', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						document.add_heading('With H Lord in same house during his dasha', level=2)
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(three2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='1' and c_ode2[2]=='0'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H10Lord', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(three2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='1' and c_ode2[2]=='1'):
						print("11")
						print("code2[i] is")
						print(code2[i])
						document.add_heading(code2[i], level=0)
						document.add_heading('H11Lord', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(three2[i])
						except:
							paragraph = document.add_paragraph('NULL')
					if(c_ode2[1]=='1' and c_ode2[2]=='2'):
						document.add_heading(code2[i], level=0)
						document.add_heading('H12Lord', level=2)
						try:
							paragraph = document.add_paragraph(one2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(two2[i])
						except:
							paragraph = document.add_paragraph('NULL')
						try:
							paragraph = document.add_paragraph(three2[i])
						except:
							paragraph = document.add_paragraph('NULL')
	
	answer3 = input('Want to EXIT ')
	if(answer3[0]=='y' or answer3[0]=='Y'):
		break
name_doc = input('Enter name of the document: ')
name_fin = name_doc+'.docx'
document.save(name_fin)
		

		


